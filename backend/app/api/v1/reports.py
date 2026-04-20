"""
CodeSentinel — Reports API
Generate PDF and DOCX security reports from real database state.
Uses Jinja2 HTML templates + WeasyPrint for PDF, python-docx for DOCX.
"""
from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Literal

import structlog
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.core.security import get_current_user
from app.models.finding import Finding
from app.models.repository import Repository
from app.models.scan import Scan
from app.models.user import User

router = APIRouter()
log = structlog.get_logger("api.reports")


async def _gather_report_data(
    org_id: str,
    db: AsyncSession,
) -> dict:
    """Gather all data needed for any report type."""
    # Repos
    repos_r = await db.execute(
        select(Repository).where(Repository.organization_id == org_id)
    )
    repos = repos_r.scalars().all()
    repo_ids = [str(r.id) for r in repos]

    if not repo_ids:
        return {"repos": [], "scans": [], "findings": [], "summary": {}}

    # Recent scans (last 30)
    scans_r = await db.execute(
        select(Scan)
        .where(Scan.repository_id.in_(repo_ids))
        .order_by(Scan.created_at.desc())
        .limit(30)
    )
    scans = scans_r.scalars().all()

    # Open findings
    findings_r = await db.execute(
        select(Finding)
        .join(Repository, Finding.repository_id == Repository.id)
        .where(
            Repository.organization_id == org_id,
            Finding.status == "open",
            Finding.is_false_positive == False,
        )
        .order_by(Finding.severity.asc(), Finding.created_at.desc())
        .limit(200)
    )
    findings = findings_r.scalars().all()

    # Severity counts
    sev_counts = {s: 0 for s in ("critical", "high", "medium", "low", "info")}
    for f in findings:
        sev_counts[f.severity] = sev_counts.get(f.severity, 0) + 1

    total_risk = sum(
        {"critical": 40, "high": 20, "medium": 8, "low": 2, "info": 0}.get(f.severity, 0)
        for f in findings
    )
    risk_score = min(100, total_risk)

    return {
        "repos": repos,
        "scans": scans,
        "findings": findings,
        "summary": {
            "total_repos": len(repos),
            "total_scans": len(scans),
            "total_findings": len(findings),
            "risk_score": risk_score,
            "sev_counts": sev_counts,
            "secrets_found": sum(1 for f in findings if f.category == "secrets"),
            "dep_vulns": sum(1 for f in findings if f.agent_type == "dependency"),
            "fixes_available": sum(1 for f in findings if f.fix_available),
            "generated_at": datetime.now(timezone.utc).isoformat(),
        },
    }


def _render_html_report(data: dict, report_type: str) -> str:
    """Render an HTML report using inline Jinja2 template."""
    from jinja2 import Environment, BaseLoader
    import html as html_lib

    findings = data["findings"]
    summary = data["summary"]
    repos = data["repos"]

    SEV_COLOR = {
        "critical": "#dc2626", "high": "#ea580c",
        "medium": "#d97706", "low": "#2563eb", "info": "#6b7280",
    }

    findings_html = ""
    for f in findings[:50]:  # Limit to 50 in PDF
        color = SEV_COLOR.get(f.severity, "#6b7280")
        findings_html += f"""
        <div style="margin-bottom:16px;padding:12px;border-left:4px solid {color};background:#f9fafb;border-radius:4px">
          <div style="display:flex;align-items:center;gap:8px;margin-bottom:6px">
            <span style="background:{color};color:white;padding:2px 8px;border-radius:12px;font-size:11px;font-weight:700">{f.severity.upper()}</span>
            <strong style="font-size:14px">{html_lib.escape(f.title)}</strong>
          </div>
          {"<div style='font-family:monospace;font-size:11px;color:#6b7280;margin-bottom:4px'>" + html_lib.escape(f.file_path or '') + (":" + str(f.line_start) if f.line_start else "") + "</div>" if f.file_path else ""}
          <div style="font-size:12px;color:#374151;margin-bottom:4px">{html_lib.escape((f.why_flagged or f.description or '')[:300])}</div>
          {"<div style='font-size:11px;color:#047857;background:#d1fae5;padding:4px 8px;border-radius:4px;margin-top:6px'><strong>Fix:</strong> " + html_lib.escape((f.recommendation or '')[:200]) + "</div>" if f.recommendation else ""}
        </div>"""

    repos_html = "".join(
        f"<tr><td>{html_lib.escape(r.full_name)}</td><td>{r.total_scans}</td>"
        f"<td>{r.open_findings}</td><td>{r.last_scan_risk_score or '—'}</td></tr>"
        for r in repos
    )

    return f"""<!DOCTYPE html><html><head>
<meta charset="UTF-8">
<style>
  body {{ font-family: -apple-system, Arial, sans-serif; color: #111827; margin: 40px; line-height: 1.6; }}
  h1 {{ color: #4f46e5; }} h2 {{ color: #1f2937; border-bottom: 2px solid #e5e7eb; padding-bottom: 8px; }}
  .stat {{ display: inline-block; margin: 8px; padding: 16px 24px; background: #f3f4f6; border-radius: 8px; text-align: center; min-width: 100px; }}
  .stat-value {{ font-size: 32px; font-weight: 900; color: #4f46e5; }}
  .stat-label {{ font-size: 12px; color: #6b7280; margin-top: 4px; }}
  table {{ width: 100%; border-collapse: collapse; margin: 16px 0; }}
  th {{ background: #f9fafb; text-align: left; padding: 8px 12px; font-size: 11px; text-transform: uppercase; color: #6b7280; border-bottom: 2px solid #e5e7eb; }}
  td {{ padding: 8px 12px; border-bottom: 1px solid #f3f4f6; font-size: 13px; }}
  .footer {{ margin-top: 40px; padding-top: 16px; border-top: 1px solid #e5e7eb; font-size: 11px; color: #9ca3af; }}
</style>
</head><body>
<h1>🛡️ CodeSentinel — {report_type.replace('_', ' ').title()}</h1>
<p style="color:#6b7280">Generated: {summary['generated_at']} · Organization report</p>

<h2>Executive Summary</h2>
<div>
  <div class="stat"><div class="stat-value" style="color:{'#dc2626' if summary['risk_score']>=80 else '#d97706' if summary['risk_score']>=40 else '#16a34a'}">{summary['risk_score']}</div><div class="stat-label">Risk Score</div></div>
  <div class="stat"><div class="stat-value">{summary['total_findings']}</div><div class="stat-label">Open Findings</div></div>
  <div class="stat"><div class="stat-value" style="color:#dc2626">{summary['sev_counts']['critical']}</div><div class="stat-label">Critical</div></div>
  <div class="stat"><div class="stat-value" style="color:#ea580c">{summary['sev_counts']['high']}</div><div class="stat-label">High</div></div>
  <div class="stat"><div class="stat-value" style="color:#047857">{summary['fixes_available']}</div><div class="stat-label">Auto-Fixes</div></div>
  <div class="stat"><div class="stat-value">{summary['total_repos']}</div><div class="stat-label">Repositories</div></div>
</div>

<h2>Repository Overview</h2>
<table><thead><tr><th>Repository</th><th>Scans</th><th>Open Findings</th><th>Last Risk Score</th></tr></thead>
<tbody>{repos_html}</tbody></table>

<h2>Findings Detail ({min(len(data['findings']), 50)} shown)</h2>
{findings_html}

<div class="footer">Generated by CodeSentinel v1.0.0 · {summary['generated_at']}</div>
</body></html>"""


@router.get("/reports/{report_type}")
async def generate_report(
    report_type: str,
    format: Literal["pdf", "docx", "json"] = Query(default="pdf"),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Generate a security report in the requested format."""
    valid_types = {"security_summary", "compliance_report", "dependency_report", "executive_summary", "findings_detail"}
    if report_type not in valid_types:
        raise HTTPException(status_code=400, detail=f"Unknown report type. Valid: {valid_types}")

    if not current_user.primary_org_id:
        raise HTTPException(status_code=400, detail="No organization found.")

    data = await _gather_report_data(current_user.primary_org_id, db)

    if format == "json":
        import json as _json
        content = _json.dumps({
            "report_type": report_type,
            "generated_at": data["summary"]["generated_at"],
            "summary": data["summary"],
            "findings": [
                {
                    "id": f.id, "title": f.title, "severity": f.severity,
                    "file_path": f.file_path, "line_start": f.line_start,
                    "description": f.description, "why_flagged": f.why_flagged,
                    "recommendation": f.recommendation, "category": f.category,
                    "agent_type": f.agent_type, "cve_id": f.cve_id,
                    "fix_available": f.fix_available,
                }
                for f in data["findings"]
            ],
        }, indent=2)
        return Response(
            content=content,
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="{report_type}.json"'},
        )

    html_content = _render_html_report(data, report_type)

    if format == "pdf":
        try:
            from weasyprint import HTML as WeasyHTML
            pdf_bytes = WeasyHTML(string=html_content).write_pdf()
            return Response(
                content=pdf_bytes,
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{report_type}.pdf"'},
            )
        except ImportError:
            raise HTTPException(
                status_code=501,
                detail="PDF generation requires WeasyPrint. Install with: pip install weasyprint",
            )
        except Exception as exc:
            log.error("PDF generation failed", error=str(exc))
            raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(exc)[:200]}")

    if format == "docx":
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches

            doc = Document()
            doc.add_heading("CodeSentinel Security Report", 0)
            doc.add_paragraph(f"Generated: {data['summary']['generated_at']}")

            doc.add_heading("Summary", level=1)
            summary = data["summary"]
            t = doc.add_table(rows=1, cols=4)
            t.style = "Table Grid"
            hdr = t.rows[0].cells
            for i, h in enumerate(["Risk Score", "Total Findings", "Critical", "High"]):
                hdr[i].text = h
            row = t.add_row().cells
            row[0].text = str(summary["risk_score"])
            row[1].text = str(summary["total_findings"])
            row[2].text = str(summary["sev_counts"]["critical"])
            row[3].text = str(summary["sev_counts"]["high"])

            doc.add_heading("Findings", level=1)
            for f in data["findings"][:50]:
                doc.add_heading(f"[{f.severity.upper()}] {f.title}", level=2)
                if f.file_path:
                    doc.add_paragraph(f"Location: {f.file_path}:{f.line_start or '?'}")
                if f.why_flagged:
                    p = doc.add_paragraph()
                    p.add_run("Why flagged: ").bold = True
                    p.add_run(f.why_flagged[:400])
                if f.recommendation:
                    p = doc.add_paragraph()
                    p.add_run("Fix: ").bold = True
                    p.add_run(f.recommendation[:300])
                doc.add_paragraph("─" * 50)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return Response(
                content=buf.read(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{report_type}.docx"'},
            )
        except ImportError:
            raise HTTPException(
                status_code=501,
                detail="DOCX generation requires python-docx. Install with: pip install python-docx",
            )
        except Exception as exc:
            log.error("DOCX generation failed", error=str(exc))
            raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(exc)[:200]}")

    raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")
